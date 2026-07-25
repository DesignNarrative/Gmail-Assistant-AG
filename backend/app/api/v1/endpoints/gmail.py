from fastapi import APIRouter, Depends, HTTPException, BackgroundTasks
from fastapi.responses import FileResponse, StreamingResponse
from sqlalchemy.ext.asyncio import AsyncSession
from sqlalchemy import select, func, desc, update
from app.core.database import get_db
from app.core.dependencies import get_current_active_user
from app.models.user import User
from app.models.email import Email
from app.models.thread import Thread
from app.models.attachment import Attachment
from app.models.sync_log import SyncLog
from app.workers.sync_tasks import run_sync_gmail_label
from pydantic import BaseModel, Field
import uuid
import logging
import io
import csv
import docx
import zipfile
import re

router = APIRouter()
logger = logging.getLogger(__name__)

class LabelSettingsUpdate(BaseModel):
    gmail_label: str = Field(..., min_length=1, max_length=100)

@router.post("/sync/trigger")
async def trigger_gmail_sync(
    background_tasks: BackgroundTasks,
    current_user: User = Depends(get_current_active_user),
    db: AsyncSession = Depends(get_db)
):
    if not current_user.google_access_token:
        raise HTTPException(
            status_code=400, 
            detail="Gmail sync is not configured. Please authorize Google OAuth first."
        )
        
    try:
        # Create a new sync log entry
        sync_log = SyncLog(
            user_id=current_user.id,
            sync_type="manual",
            status="running"
        )
        db.add(sync_log)
        await db.commit()
        await db.refresh(sync_log)
        
        # Trigger background task directly inside FastAPI's event loop
        background_tasks.add_task(run_sync_gmail_label, str(current_user.id), str(sync_log.id))
        
        return {
            "detail": "Gmail sync job successfully started in background.",
            "sync_job_id": sync_log.id,
            "status": sync_log.status,
            "started_at": sync_log.started_at
        }
    except Exception as e:
        logger.error(f"Failed to trigger sync: {e}")
        raise HTTPException(status_code=500, detail="Failed to schedule sync job")

@router.get("/sync/status")
async def get_sync_status(
    limit: int = 10,
    current_user: User = Depends(get_current_active_user),
    db: AsyncSession = Depends(get_db)
):
    try:
        stmt = (
            select(SyncLog)
            .where(SyncLog.user_id == current_user.id)
            .order_by(desc(SyncLog.started_at))
            .limit(limit)
        )
        result = await db.execute(stmt)
        logs = result.scalars().all()
        return logs
    except Exception as e:
        logger.error(f"Error fetching sync status: {e}")
        raise HTTPException(status_code=500, detail="Failed to retrieve sync logs")

@router.get("/sync/stats")
async def get_sync_stats(
    current_user: User = Depends(get_current_active_user),
    db: AsyncSession = Depends(get_db)
):
    try:
        # Total emails for this user
        emails_count_stmt = select(func.count(Email.id)).where(Email.user_id == current_user.id)
        total_emails = (await db.execute(emails_count_stmt)).scalar() or 0
        
        # Total distinct threads for this user
        threads_count_stmt = select(func.count(func.distinct(Email.thread_id))).where(Email.user_id == current_user.id)
        total_threads = (await db.execute(threads_count_stmt)).scalar() or 0
        
        # Total attachments belonging to this user's emails
        att_count_stmt = (
            select(func.count(Attachment.id))
            .join(Email, Email.id == Attachment.email_id)
            .where(Email.user_id == current_user.id)
        )
        total_attachments = (await db.execute(att_count_stmt)).scalar() or 0
        
        # Total attachment size belonging to this user's emails
        att_size_stmt = (
            select(func.sum(Attachment.file_size))
            .join(Email, Email.id == Attachment.email_id)
            .where(Email.user_id == current_user.id)
        )
        total_size_bytes = (await db.execute(att_size_stmt)).scalar() or 0
        
        # Fetch latest sync run for this user
        latest_sync_stmt = (
            select(SyncLog)
            .where(SyncLog.user_id == current_user.id)
            .order_by(desc(SyncLog.started_at))
            .limit(1)
        )
        latest_sync = (await db.execute(latest_sync_stmt)).scalars().first()
        
        return {
            "total_emails": total_emails,
            "total_threads": total_threads,
            "total_attachments": total_attachments,
            "total_size_bytes": total_size_bytes,
            "latest_sync": latest_sync
        }
    except Exception as e:
        logger.error(f"Error fetching sync stats: {e}")
        raise HTTPException(status_code=500, detail="Failed to retrieve sync statistics")

@router.put("/settings/label")
async def update_gmail_label(
    payload: LabelSettingsUpdate,
    current_user: User = Depends(get_current_active_user),
    db: AsyncSession = Depends(get_db)
):
    try:
        stmt = update(User).where(User.id == current_user.id).values(gmail_label=payload.gmail_label)
        await db.execute(stmt)
        await db.commit()
        return {"detail": "Gmail sync label updated successfully.", "gmail_label": payload.gmail_label}
    except Exception as e:
        logger.error(f"Error updating Gmail sync label: {e}")
        raise HTTPException(status_code=500, detail="Failed to update sync label settings")


@router.get("/export")
async def export_synced_emails(
    download_all: bool = False,
    current_user: User = Depends(get_current_active_user),
    db: AsyncSession = Depends(get_db)
):
    try:
        from fastapi import Response
        # By default ("Download New") export only emails not yet downloaded.
        # When download_all=true ("Download All") export every synced email regardless of flag.
        conditions = [Email.user_id == current_user.id]
        if not download_all:
            conditions.append(Email.is_downloaded == False)
        stmt = select(Email).where(*conditions).order_by(desc(Email.date_sent))
        result = await db.execute(stmt)
        emails = result.scalars().all()

        if not emails:
            # Return 204 No Content if everything is already downloaded
            return Response(status_code=204)

        zip_buffer = io.BytesIO()
        with zipfile.ZipFile(zip_buffer, "w", zipfile.ZIP_DEFLATED) as zip_file:
            for idx, e in enumerate(emails):
                doc = docx.Document()
                
                # Title Heading
                h = doc.add_heading(e.subject or "(No Subject)", level=1)
                h.keep_with_next = True
                
                # Metadata block
                date_sent_str = e.date_sent.strftime("%Y-%m-%d %H:%M:%S") if e.date_sent else "(No Date)"
                doc.add_paragraph(
                    f"From: {e.sender_name or ''} <{e.sender_email}>\n"
                    f"Date: {date_sent_str}\n"
                    f"Thread ID: {e.thread_id}"
                )
                doc.add_paragraph("-" * 60)
                
                # Clean up raw email body formatting and zero-width spaces/junk characters
                body_raw = e.body_text or "(No body content)"
                body_clean = re.sub(r'[\u200b-\u200d\ufeff\u034f\u200e\u200f]', '', body_raw)
                body_clean = re.sub(r'\n{3,}', '\n\n', body_clean)
                body_clean = re.sub(r'[ \t]{2,}', ' ', body_clean).strip()
                
                # Write body text lines to document
                for line in body_clean.split("\n"):
                    doc.add_paragraph(line)
                
                # Save single docx file to memory
                doc_buf = io.BytesIO()
                doc.save(doc_buf)
                docx_bytes = doc_buf.getvalue()
                
                # Derive safe file name
                safe_subject = re.sub(r'[^\w\s-]', '', e.subject or "No_Subject")
                safe_subject = re.sub(r'[-\s]+', '_', safe_subject).strip('_')
                if not safe_subject:
                    safe_subject = f"email_{idx}"
                filename = f"{safe_subject[:60]}.docx"
                
                # Add file to ZIP archive
                zip_file.writestr(filename, docx_bytes)
                
                # Mark as downloaded in DB
                e.is_downloaded = True

        await db.commit()
        zip_buffer.seek(0)

        zip_name = "all_emails.zip" if download_all else "new_emails.zip"
        return StreamingResponse(
            zip_buffer,
            media_type="application/zip",
            headers={"Content-Disposition": f"attachment; filename={zip_name}"}
        )
    except Exception as e:
        logger.error(f"Error exporting synced emails to zip: {e}", exc_info=True)
        raise HTTPException(status_code=500, detail="Failed to export synced emails")


@router.get("/attachments/{attachment_id}/download")
async def download_attachment(
    attachment_id: str,
    current_user: User = Depends(get_current_active_user),
    db: AsyncSession = Depends(get_db)
):
    try:
        att_uuid = uuid.UUID(attachment_id)
    except ValueError:
        raise HTTPException(status_code=400, detail="Invalid attachment ID format")

    try:
        # Join Email to verify the attachment belongs to an email owned by current_user.id
        stmt = (
            select(Attachment, Email)
            .join(Email, Email.id == Attachment.email_id)
            .where(Attachment.id == att_uuid, Email.user_id == current_user.id)
        )
        row = (await db.execute(stmt)).first()

        if not row:
            raise HTTPException(status_code=404, detail="Attachment not found or access denied")

        attachment, email = row

        import os
        if not os.path.exists(attachment.storage_path):
            raise HTTPException(status_code=404, detail="Attachment file not found on disk")

        return FileResponse(
            path=attachment.storage_path,
            filename=attachment.filename,
            media_type=attachment.mime_type
        )
    except HTTPException as he:
        raise he
    except Exception as e:
        logger.error(f"Error downloading attachment {attachment_id}: {e}")
        raise HTTPException(status_code=500, detail="Failed to download attachment")

